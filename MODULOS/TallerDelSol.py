import streamlit as st
import pandas as pd
from fpdf import FPDF
from docx import Document


st.title("Taller del Sol")
def app():
    st.title("Generador de Remisiones")

# Cargar datos de proveedores y productos desde archivos Excel
proveedores_df = pd.read_excel("proveedores.xlsx")  # Asegúrate de tener este archivo
productos_df = pd.read_excel("productos.xlsx")      # Asegúrate de tener este archivo

# Función para obtener la información del proveedor
def buscar_proveedor(nombre_proveedor):
    return proveedores_df[proveedores_df["nombre"] == nombre_proveedor].iloc[0]

# Función para obtener la información del producto por ISBN
def buscar_producto(isbn):
    return productos_df[productos_df["ISBN"] == isbn].iloc[0]

# Crear la estructura del archivo PDF
def generar_pdf(remision_data, total, total_con_descuento):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 12)
    pdf.cell(200, 10, "Remision", ln=True, align="C")
    
    # Datos del proveedor
    pdf.set_font("Arial", size=10)
    for key, value in remision_data["proveedor"].items():
        pdf.cell(200, 10, f"{key}: {value}", ln=True)

    pdf.ln(10)

    # Detalles de productos
    pdf.cell(50, 10, "ISBN", 1)
    pdf.cell(50, 10, "Nombre", 1)
    pdf.cell(50, 10, "Cantidad", 1)
    pdf.cell(50, 10, "Precio", 1)
    pdf.ln()

    for producto in remision_data["productos"]:
        pdf.cell(50, 10, producto["ISBN"], 1)
        pdf.cell(50, 10, producto["nombre"], 1)
        pdf.cell(50, 10, str(producto["cantidad"]), 1)
        pdf.cell(50, 10, f"${producto['precio']:.2f}", 1)
        pdf.ln()

    pdf.ln(10)
    pdf.cell(200, 10, f"Total: ${total:.2f}", ln=True)
    pdf.cell(200, 10, f"Total con descuento: ${total_con_descuento:.2f}", ln=True)

    pdf.output("remision.pdf")

# Crear la estructura del archivo Word
def generar_word(remision_data, total, total_con_descuento):
    doc = Document()
    doc.add_heading("Remision", 0)

    # Datos del proveedor
    for key, value in remision_data["proveedor"].items():
        doc.add_paragraph(f"{key}: {value}")

    doc.add_paragraph("\nDetalles de productos:")
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ISBN'
    hdr_cells[1].text = 'Nombre'
    hdr_cells[2].text = 'Cantidad'
    hdr_cells[3].text = 'Precio'

    for producto in remision_data["productos"]:
        row_cells = table.add_row().cells
        row_cells[0].text = producto["ISBN"]
        row_cells[1].text = producto["nombre"]
        row_cells[2].text = str(producto["cantidad"])
        row_cells[3].text = f"${producto['precio']:.2f}"

    doc.add_paragraph(f"\nTotal: ${total:.2f}")
    doc.add_paragraph(f"Total con descuento: ${total_con_descuento:.2f}")

    doc.save("remision.docx")

# Inputs del usuario
proveedor_input = st.text_input("Nombre del Proveedor")
isbn_input = st.text_input("ISBN del Producto")
cantidad_input = st.number_input("Cantidad", min_value=1, step=1)
descuento = st.number_input("Descuento (%)", min_value=0.0, max_value=100.0)

# Lista para almacenar los productos agregados
productos = []

# Botón para agregar el producto a la lista
if st.button("Agregar Producto"):
    if isbn_input and cantidad_input:
        try:
            producto_data = buscar_producto(isbn_input.strip())
            productos.append({
                "ISBN": isbn_input.strip(),
                "nombre": producto_data["nombre"],
                "cantidad": cantidad_input,
                "precio": producto_data["precio"] * cantidad_input
            })
            st.success(f"Producto {isbn_input} agregado.")
        except IndexError:
            st.error("Producto no encontrado.")

# Mostrar productos agregados
st.write("Productos agregados:")
for prod in productos:
    st.write(f"ISBN: {prod['ISBN']}, Nombre: {prod['nombre']}, Cantidad: {prod['cantidad']}, Precio: ${prod['precio']:.2f}")

if st.button("Generar Remisión"):
    # Buscar información del proveedor
    try:
        proveedor_data = buscar_proveedor(proveedor_input)
    except IndexError:
        st.error("Proveedor no encontrado.")
        st.stop()

    # Calcular el total y el total con descuento
    total = sum([prod["precio"] for prod in productos])
    total_con_descuento = total * (1 - descuento / 100)

    # Crear remisión
    remision_data = {
        "proveedor": proveedor_data.to_dict(),
        "productos": productos
    }

    # Generar archivos PDF y Word
    generar_pdf(remision_data, total, total_con_descuento)
    generar_word(remision_data, total, total_con_descuento)

    # Mostrar resultados
    st.success("Remisión generada con éxito.")
    st.download_button("Descargar PDF", data=open("remision.pdf", "rb"), file_name="remision.pdf")
    st.download_button("Descargar Word", data=open("remision.docx", "rb"), file_name="remision.docx")
